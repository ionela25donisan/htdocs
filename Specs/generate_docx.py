from docx import Document
from pathlib import Path

path = Path('DocumentWord_SiteGestionare_Apartamente.md')
text = path.read_text(encoding='utf-8')
lines = text.splitlines()
doc = Document()
for line in lines:
    if line.startswith('### '):
        doc.add_heading(line[4:], level=3)
        continue
    if line.startswith('## '):
        doc.add_heading(line[3:], level=2)
        continue
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
        continue
    if line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
        continue
    if line.strip() == '---':
        doc.add_page_break()
        continue
    if line.strip() == '':
        doc.add_paragraph('')
        continue
    if line.startswith('```'):
        continue
    doc.add_paragraph(line)

out = Path('DocumentWord_SiteGestionare_Apartamente.docx')
doc.save(out)
print('Created', out)
